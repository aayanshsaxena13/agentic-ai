# Import packages here for use...
from docx import Document
import os
import json

# Create new helpers...
file_path = os.path.join(os.path.dirname(__file__), "..", "input", "word_input.json") # or give your own input...{ header:str, paragraphs:list of str }
output_path = os.path.join(os.path.dirname(__file__), "..", "output", "word.docx") # change to your own output path

def createWordFile():
    doc = Document() # create a new word document

    # Read input file...
    with open(file_path, "r") as f:
        input_file = json.loads(f.read())

    # Write some things...
    doc.add_heading(input_file["header"], level=6)

    for para in input_file["paragraphs"]:
        doc.add_paragraph(para)

    doc.save(output_path)
    return "Generated word document using Agentic AI..."

def readWordFile(path):
    doc = Document(path)
    return doc.paragraphs